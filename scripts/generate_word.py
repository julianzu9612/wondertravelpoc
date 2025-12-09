from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# Configuration
OUTPUT_FILE = "Propuesta_Comercial_Wonder_Travel.docx"
LOGO_PATH = os.path.abspath("ORBITAL LAB LOGO.png")

def create_word_doc():
    doc = Document()

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- Header with Logo ---
    # We'll put the logo in the header or top of body. Top of body is easier for alignment control.
    if os.path.exists(LOGO_PATH):
        try:
            doc.add_picture(LOGO_PATH, width=Inches(2.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Warning: Could not add logo: {e}")
    
    doc.add_paragraph("") # Spacer

    # --- Title ---
    title = doc.add_heading('Propuesta de Modernización Tecnológica & Eficiencia Operativa', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- Metadata ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run("Para: Juan Pablo Gaviria, CEO Wonder Travel\n")
    runner.bold = True
    runner = p.add_run("De: Orbital Lab\n")
    runner.bold = True
    runner = p.add_run("Fecha: 2 de Diciembre de 2025")

    doc.add_paragraph("---").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 1. Executive Summary ---
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    p = doc.add_paragraph(
        "Wonder Travel se encuentra en un punto de inflexión. Su marca y clientela son de clase mundial (Harvard, MIT), "
        "pero su infraestructura tecnológica actual opera con ineficiencias críticas que generan sobrecostos del "
    )
    p.add_run("~900%").bold = True
    p.add_run(" y limitan su agilidad comercial.")
    
    p = doc.add_paragraph(
        "Esta propuesta presenta una estrategia de "
    )
    p.add_run('"Reconstrucción Inteligente"').bold = True
    p.add_run(
        " para migrar su plataforma a una arquitectura moderna, reduciendo drásticamente los costos fijos y "
        "elevando la experiencia de usuario al nivel Premium que su marca exige."
    )

    # --- 2. Diagnosis ---
    doc.add_heading('2. Diagnóstico Financiero y Técnico', level=1)
    doc.add_paragraph("Hemos auditado su infraestructura actual y detectado una desconexión entre el gasto y el valor recibido:")

    # --- TABLE ---
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = False 
    table.allow_autofit = False
    
    # Set column widths
    for cell in table.columns[0].cells: cell.width = Inches(1.2)
    for cell in table.columns[1].cells: cell.width = Inches(2.0)
    for cell in table.columns[2].cells: cell.width = Inches(2.0)
    for cell in table.columns[3].cells: cell.width = Inches(1.3)

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Concepto'
    hdr_cells[1].text = 'Situación Actual\n(Google Cloud)'
    hdr_cells[2].text = 'Situación Propuesta\n(Modern Stack)'
    hdr_cells[3].text = 'Impacto'

    # Style Header
    for cell in hdr_cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcVAlign = parse_xml(r'<w:vAlign %s w:val="center"/>' % nsdecls('w'))
        tcPr.append(tcVAlign)
        shading_elm = parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w')))
        tcPr.append(shading_elm)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    # Data Rows
    data = [
        ("Costo Mensual", "$406.19 USD", "~$45.00 USD", "90% Ahorro"),
        ("Costo Anual", "~$4,874 USD", "~$540 USD", "+$4,300 USD\nUtilidad"),
        ("Tecnología", "Legacy\n(Lenta, Compleja)", "Next.js 14\n(Instantánea)", "Mejor UX/SEO"),
        ("Mantenimiento", "Requiere Ingenieros Especializados", "Automatizado /\nServerless", "Menor Dependencia")
    ]

    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        
        # Bold the first column
        row_cells[0].paragraphs[0].runs[0].bold = True
        
        # Center align columns 1, 2, 3
        for i in range(1, 4):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("") # Spacer

    quote = doc.add_paragraph()
    quote.add_run("Conclusión: ").bold = True
    quote.add_run('Actualmente, Wonder Travel paga un "impuesto tecnológico" de ')
    quote.add_run("~$4,300 USD anuales").bold = True
    quote.add_run(" por recursos de servidor que no utiliza ni necesita.")
    quote.style = 'Quote'

    # --- 3. Solution ---
    doc.add_heading('3. Nuestra Solución: Transformación Integral', level=1)
    doc.add_paragraph('No proponemos un simple "parche", sino una modernización completa en 4 semanas:')

    doc.add_heading('A. Eficiencia Financiera (Backend)', level=2)
    doc.add_paragraph(
        "Eliminaremos la compleja infraestructura de Google Cloud. Migraremos su base de datos y lógica de negocio "
        "a una arquitectura Serverless (Supabase + Vercel), que cobra por uso real, no por servidor encendido."
    )

    doc.add_heading('B. Excelencia Visual (Frontend)', level=2)
    doc.add_paragraph(
        "Aprovecharemos la migración para renovar la interfaz de usuario. Implementaremos un Diseño Minimalista Premium "
        "(estilo Airbnb/Apple), optimizado para móviles, con tiempos de carga instantáneos que mejorarán la conversión "
        "de sus clientes corporativos."
    )

    # --- 4. Investment ---
    doc.add_heading('4. Inversión y Retorno (ROI)', level=1)
    doc.add_paragraph("Hemos diseñado una estructura de inversión simplificada para maximizar su retorno a corto plazo.")

    doc.add_heading('Inversión Única del Proyecto', level=2)
    p = doc.add_paragraph()
    p.add_run("$1,900 USD").bold = True
    p.add_run(" (~$8,000,000 COP)")

    doc.add_paragraph("Incluye: Migración completa de datos, desarrollo de nueva plataforma (Frontend + Backend), configuración de servidores, pruebas y puesta en producción.", style='List Bullet')
    doc.add_paragraph("Tiempo de Ejecución: 4 Semanas.", style='List Bullet')

    doc.add_heading('Retorno de Inversión (ROI)', level=2)
    p = doc.add_paragraph(
        "Gracias al ahorro directo en costos de hosting ($360 USD/mes de ahorro), "
    )
    p.add_run("este proyecto se paga solo en 5.2 meses").bold = True
    p.add_run(".")
    doc.add_paragraph("A partir del mes 6, la modernización genera flujo de caja positivo directo a la utilidad de la empresa.")

    # --- 5. Next Steps ---
    doc.add_heading('5. Próximos Pasos', level=1)
    doc.add_paragraph("Para iniciar esta transformación y detener el desperdicio de recursos:")
    
    doc.add_paragraph("Aprobación: Firma de acuerdo de servicios.", style='List Number')
    doc.add_paragraph("Acceso: Entrega de backup de base de datos actual.", style='List Number')
    doc.add_paragraph("Ejecución: Inicio inmediato del desarrollo del \"Gemelo Digital\" (sin afectar su sitio actual).", style='List Number')

    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer = doc.add_paragraph("En Orbital Lab, no vendemos código; construimos activos digitales eficientes que impulsan su negocio.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style = 'Subtitle'

    doc.save(OUTPUT_FILE)
    print(f"Word document generated: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_word_doc()
